import sys
import os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

base = r"[PUB] India_runs_data_and_ai_challenge\India_runs_data_and_ai_challenge"
files = ["job_description.docx", "redrob_signals_doc.docx", "submission_spec.docx"]

output = []
for f in files:
    path = os.path.join(base, f)
    output.append(f"\n{'='*80}")
    output.append(f"=== {f} ===")
    output.append(f"{'='*80}")
    doc = Document(path)
    for para in doc.paragraphs:
        if para.text.strip():
            output.append(para.text)
    for i, table in enumerate(doc.tables):
        output.append(f"\n--- Table {i+1} ---")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            output.append(" | ".join(cells))

with open("docx_contents.txt", "w", encoding="utf-8") as fout:
    fout.write("\n".join(output))

print(f"Written {len(output)} lines to docx_contents.txt")
